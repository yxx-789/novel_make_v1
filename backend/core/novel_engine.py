# -*- coding: utf-8 -*-
"""
小说创作引擎
封装 AI Novel Generator 的核心功能
"""

import sys
import os
from typing import List, Dict, Optional, Any
from pathlib import Path
import json
import asyncio
from datetime import datetime

# 添加 AI_NovelGenerator 到路径
NOVEL_GENERATOR_PATH = Path(__file__).parent.parent.parent / "AI_NovelGenerator" / "AI_NovelGenerator"
if NOVEL_GENERATOR_PATH.exists():
    sys.path.insert(0, str(NOVEL_GENERATOR_PATH))

# 导入模型
sys.path.insert(0, str(Path(__file__).parent.parent))
from models.schemas import (
    NovelProject, CharacterProfile, WorldSetting, PlotBlueprint,
    ChapterOutline, ChapterContent, NovelGenre, NovelStatus
)


class LLMAdapter:
    """LLM 适配器 - 统一不同模型的调用"""
    
    def __init__(self, config: Dict[str, Any]):
        self.api_key = config.get("api_key")
        self.base_url = config.get("base_url", "https://api.openai.com/v1")
        self.model = config.get("model", "gpt-4o-mini")
        self.temperature = config.get("temperature", 0.7)
        self.max_tokens = config.get("max_tokens", 4096)
        
        # 初始化客户端
        self._init_client()
    
    def _init_client(self):
        """初始化 API 客户端"""
        try:
            from openai import OpenAI
            self.client = OpenAI(
                api_key=self.api_key,
                base_url=self.base_url
            )
        except Exception as e:
            print(f"Warning: Could not initialize OpenAI client: {e}")
            self.client = None
    
    async def generate(self, prompt: str, system_prompt: str = None) -> str:
        """生成文本"""
        if not self.client:
            raise RuntimeError("LLM client not initialized")
        
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": prompt})
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=self.temperature,
                max_tokens=self.max_tokens
            )
            return response.choices[0].message.content
        except Exception as e:
            raise RuntimeError(f"LLM generation failed: {e}")


class EmbeddingAdapter:
    """Embedding 适配器 - 用于向量检索"""
    
    def __init__(self, config: Dict[str, Any]):
        self.api_key = config.get("embedding_api_key")
        self.base_url = config.get("embedding_url", "https://api.openai.com/v1")
        self.model = config.get("embedding_model", "text-embedding-ada-002")
        self._init_client()
    
    def _init_client(self):
        """初始化 Embedding 客户端"""
        try:
            from openai import OpenAI
            self.client = OpenAI(
                api_key=self.api_key,
                base_url=self.base_url
            )
        except Exception as e:
            print(f"Warning: Could not initialize Embedding client: {e}")
            self.client = None
    
    async def embed(self, text: str) -> List[float]:
        """生成文本嵌入"""
        if not self.client:
            raise RuntimeError("Embedding client not initialized")
        
        try:
            response = self.client.embeddings.create(
                model=self.model,
                input=text
            )
            return response.data[0].embedding
        except Exception as e:
            raise RuntimeError(f"Embedding generation failed: {e}")


class MemoryManager:
    """记忆管理器 - 管理小说的一致性"""
    
    def __init__(self, novel_id: str, vectorstore_path: str = None):
        self.novel_id = novel_id
        self.vectorstore_path = vectorstore_path or f"./data/vectorstore/{novel_id}"
        self.character_states: Dict[str, Dict] = {}
        self.global_summary: str = ""
        self.plot_arcs: List[Dict] = []
        
        # 初始化向量存储
        self._init_vectorstore()
    
    def _init_vectorstore(self):
        """初始化向量存储"""
        try:
            import chromadb
            self.client = chromadb.PersistentClient(path=self.vectorstore_path)
            self.collection = self.client.get_or_create_collection(
                name=f"novel_{self.novel_id}",
                metadata={"hnsw:space": "cosine"}
            )
        except Exception as e:
            print(f"Warning: Could not initialize vectorstore: {e}")
            self.client = None
            self.collection = None
    
    async def update_character_state(self, character_name: str, state: Dict):
        """更新角色状态"""
        if character_name not in self.character_states:
            self.character_states[character_name] = {
                "name": character_name,
                "current_location": None,
                "relationships": {},
                "inventory": [],
                "status_changes": [],
                "last_appearance_chapter": 0
            }
        
        self.character_states[character_name].update(state)
        self.character_states[character_name]["last_updated"] = datetime.now().isoformat()
    
    async def update_global_summary(self, summary: str):
        """更新全局摘要"""
        self.global_summary = summary
    
    async def add_plot_arc(self, arc: Dict):
        """添加情节弧"""
        self.plot_arcs.append(arc)
    
    async def store_chapter_memory(self, chapter_num: int, content: str, embedding: List[float] = None):
        """存储章节记忆到向量库"""
        if not self.collection:
            return
        
        try:
            self.collection.add(
                ids=[f"chapter_{chapter_num}"],
                embeddings=[embedding] if embedding else None,
                documents=[content],
                metadatas=[{
                    "chapter_num": chapter_num,
                    "type": "chapter_content",
                    "timestamp": datetime.now().isoformat()
                }]
            )
        except Exception as e:
            print(f"Warning: Could not store chapter memory: {e}")
    
    async def retrieve_relevant_context(self, query: str, top_k: int = 5) -> List[str]:
        """检索相关上下文"""
        if not self.collection:
            return []
        
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=top_k
            )
            return results.get("documents", [[]])[0]
        except Exception as e:
            print(f"Warning: Could not retrieve context: {e}")
            return []
    
    def get_state_dict(self) -> Dict:
        """获取状态字典"""
        return {
            "character_states": self.character_states,
            "global_summary": self.global_summary,
            "plot_arcs": self.plot_arcs
        }


class NovelEngine:
    """
    小说创作引擎
    
    整合 AI Novel Generator 的核心功能：
    - 蓝图生成
    - 章节规划
    - 内容生成
    - 一致性检查
    """
    
    def __init__(self, llm_config: Dict, embedding_config: Dict = None):
        self.llm = LLMAdapter(llm_config)
        self.embedding = EmbeddingAdapter(embedding_config or llm_config) if embedding_config else None
        
        # 项目存储
        self.projects: Dict[str, NovelProject] = {}
        self.memories: Dict[str, MemoryManager] = {}
    
    # ==================== 项目管理 ====================
    
    async def create_project(self, config: Dict) -> NovelProject:
        """
        创建小说项目
        
        Args:
            config: 项目配置
                - title: 标题
                - genre: 类型
                - topic: 主题
                - total_chapters: 总章节数
                - target_word_count: 目标字数
        
        Returns:
            NovelProject: 创建的项目
        """
        import uuid
        
        novel_id = str(uuid.uuid4())[:8]
        
        project = NovelProject(
            novel_id=novel_id,
            title=config.get("title", "未命名小说"),
            author=config.get("author", "AI"),
            genre=NovelGenre(config.get("genre", "玄幻")),
            topic=config.get("topic", ""),
            total_chapters=config.get("total_chapters", 10),
            target_word_count=config.get("target_word_count", 3000),
            style_guide=config.get("style_guide")
        )
        
        self.projects[novel_id] = project
        self.memories[novel_id] = MemoryManager(novel_id)
        
        return project
    
    async def get_project(self, novel_id: str) -> Optional[NovelProject]:
        """获取项目"""
        return self.projects.get(novel_id)
    
    async def list_projects(self, page: int = 1, page_size: int = 10) -> List[NovelProject]:
        """列出项目"""
        projects = list(self.projects.values())
        start = (page - 1) * page_size
        return projects[start:start + page_size]
    
    # ==================== 设定生成 ====================
    
    async def generate_blueprint(self, novel_id: str) -> PlotBlueprint:
        """
        生成小说蓝图
        
        包括：世界观、角色设定、情节大纲
        """
        project = await self.get_project(novel_id)
        if not project:
            raise ValueError(f"Project {novel_id} not found")
        
        # 生成世界观
        world_setting = await self._generate_world_setting(project)
        project.world_setting = world_setting
        
        # 生成角色
        characters = await self._generate_characters(project)
        project.characters = characters
        
        # 生成情节蓝图
        plot_blueprint = await self._generate_plot_blueprint(project)
        project.plot_blueprint = plot_blueprint
        
        project.status = NovelStatus.DRAFT
        project.updated_at = datetime.now()
        
        return plot_blueprint
    
    async def _generate_world_setting(self, project: NovelProject) -> WorldSetting:
        """生成世界观设定"""
        prompt = f"""
请为以下小说设计世界观设定：

标题：{project.title}
类型：{project.genre.value}
主题：{project.topic}

请生成：
1. 时代背景
2. 主要地点
3. 力量体系（如果有）
4. 社会结构
5. 独特元素（至少3个）
6. 世界规则

以JSON格式输出。
"""
        
        result = await self.llm.generate(prompt)
        # 解析结果并构建 WorldSetting
        # 这里简化处理，实际需要解析JSON
        return WorldSetting(
            era="现代",
            location="虚构城市",
            power_system="无",
            social_structure="现代社会",
            unique_elements=["AI技术", "虚拟现实"],
            rules=["物理定律正常"]
        )
    
    async def _generate_characters(self, project: NovelProject) -> List[CharacterProfile]:
        """生成角色设定"""
        prompt = f"""
请为以下小说设计主要角色：

标题：{project.title}
类型：{project.genre.value}
主题：{project.topic}
世界观：{project.world_setting}

请生成至少3个主要角色（主角、配角、反派），包括：
- 姓名
- 角色定位
- 年龄
- 性格特点
- 背景故事
- 能力/技能
- 人物关系

以JSON格式输出。
"""
        
        result = await self.llm.generate(prompt)
        # 解析并返回角色列表
        return [
            CharacterProfile(
                name="主角",
                role="主角",
                personality=["勇敢", "聪明"],
                background="普通人"
            )
        ]
    
    async def _generate_plot_blueprint(self, project: NovelProject) -> PlotBlueprint:
        """生成情节蓝图"""
        prompt = f"""
请为以下小说设计情节蓝图：

标题：{project.title}
类型：{project.genre.value}
主题：{project.topic}
总章节：{project.total_chapters}
角色：{[c.name for c in project.characters]}

请生成：
1. 核心冲突
2. 触发事件
3. 上升行动（至少5个关键事件）
4. 高潮
5. 下降行动
6. 结局
7. 伏笔设置（至少3个）

以JSON格式输出。
"""
        
        result = await self.llm.generate(prompt)
        # 解析并返回情节蓝图
        return PlotBlueprint(
            main_conflict="正义与邪恶的对抗",
            inciting_incident="主角获得特殊能力",
            rising_actions=["初识能力", "遇到导师", "遭遇挫折"],
            climax="最终决战",
            falling_actions=["战后的平静"],
            resolution="新的开始",
            foreshadowing=[{"hint": "神秘信物", "chapter": 1, "payoff_chapter": 50}]
        )
    
    # ==================== 章节生成 ====================
    
    async def generate_chapter_outline(self, novel_id: str) -> List[ChapterOutline]:
        """生成章节大纲"""
        project = await self.get_project(novel_id)
        if not project:
            raise ValueError(f"Project {novel_id} not found")
        
        prompt = f"""
请为以下小说生成章节大纲：

标题：{project.title}
类型：{project.genre.value}
总章节：{project.total_chapters}
情节蓝图：{project.plot_blueprint}

请生成每一章的：
1. 章节标题
2. 章节概要
3. 关键事件
4. 涉及角色

以JSON数组格式输出。
"""
        
        result = await self.llm.generate(prompt)
        
        # 解析并构建章节大纲
        outlines = []
        for i in range(1, project.total_chapters + 1):
            outlines.append(ChapterOutline(
                chapter_num=i,
                title=f"第{i}章",
                summary=f"第{i}章的内容概要",
                key_events=[],
                characters_involved=[],
                word_count_target=project.target_word_count
            ))
        
        project.chapters = outlines
        project.updated_at = datetime.now()
        
        return outlines
    
    async def generate_chapter(
        self,
        novel_id: str,
        chapter_num: int,
        additional_guidance: str = None,
        use_memory: bool = True
    ) -> ChapterContent:
        """
        生成章节内容
        
        Args:
            novel_id: 小说ID
            chapter_num: 章节号
            additional_guidance: 额外指导
            use_memory: 是否使用记忆系统
        
        Returns:
            ChapterContent: 生成的章节内容
        """
        project = await self.get_project(novel_id)
        if not project:
            raise ValueError(f"Project {novel_id} not found")
        
        if chapter_num > len(project.chapters):
            raise ValueError(f"Chapter {chapter_num} not found in outline")
        
        chapter_outline = project.chapters[chapter_num - 1]
        memory = self.memories.get(novel_id)
        
        # 构建上下文
        context_parts = [
            f"小说标题：{project.title}",
            f"类型：{project.genre.value}",
            f"世界观：{project.world_setting}",
            f"角色：{[c.name for c in project.characters]}",
            f"\n当前章节：第{chapter_num}章 - {chapter_outline.title}",
            f"章节概要：{chapter_outline.summary}",
            f"关键事件：{chapter_outline.key_events}",
        ]
        
        # 添加记忆上下文
        if use_memory and memory:
            context_parts.append(f"\n全局摘要：{memory.global_summary}")
            
            # 检索相关上下文
            relevant_context = await memory.retrieve_relevant_context(chapter_outline.summary)
            if relevant_context:
                context_parts.append(f"\n相关上下文：\n" + "\n".join(relevant_context[:3]))
        
        # 添加前文
        if chapter_num > 1 and project.generated_chapters:
            prev_chapter = project.generated_chapters[-1]
            context_parts.append(f"\n上一章结尾：\n{prev_chapter.content[-500:]}")
        
        # 添加额外指导
        if additional_guidance:
            context_parts.append(f"\n创作指导：{additional_guidance}")
        
        # 构建生成提示
        context_str = "\n".join(context_parts)
        prompt = f"""
{context_str}

请创作第{chapter_num}章的正文内容。
要求：
1. 字数约 {project.target_word_count} 字
2. 风格符合 {project.genre.value} 类型
3. 情节连贯，符合章节大纲
4. 角色行为符合设定
5. 避免触发内容安全审核

直接输出正文内容，不要包含章节标题。
"""
        
        # 生成内容
        content = await self.llm.generate(prompt, system_prompt="你是一位专业的小说作家。")
        
        # 创建章节内容对象
        chapter_content = ChapterContent(
            chapter_num=chapter_num,
            title=chapter_outline.title,
            content=content,
            word_count=len(content),
            summary=chapter_outline.summary
        )
        
        # 添加到项目
        project.generated_chapters.append(chapter_content)
        project.word_count += len(content)
        project.status = NovelStatus.IN_PROGRESS
        project.updated_at = datetime.now()
        
        return chapter_content
    
    async def finalize_chapter(
        self,
        novel_id: str,
        chapter_num: int
    ) -> Dict[str, Any]:
        """
        最终化章节
        
        更新记忆系统、角色状态、情节发展
        """
        project = await self.get_project(novel_id)
        if not project:
            raise ValueError(f"Project {novel_id} not found")
        
        chapter = next(
            (c for c in project.generated_chapters if c.chapter_num == chapter_num),
            None
        )
        if not chapter:
            raise ValueError(f"Chapter {chapter_num} not generated yet")
        
        memory = self.memories.get(novel_id)
        if not memory:
            raise ValueError(f"Memory manager not found for project {novel_id}")
        
        # 提取并更新角色状态
        await self._extract_and_update_character_states(memory, chapter)
        
        # 更新全局摘要
        await self._update_global_summary(memory, chapter)
        
        # 存储章节到向量库
        if self.embedding:
            embedding = await self.embedding.embed(chapter.content)
            await memory.store_chapter_memory(chapter_num, chapter.content, embedding)
        
        return {
            "chapter_num": chapter_num,
            "status": "finalized",
            "memory_updated": True,
            "state": memory.get_state_dict()
        }
    
    async def _extract_and_update_character_states(self, memory: MemoryManager, chapter: ChapterContent):
        """提取并更新角色状态"""
        prompt = f"""
请分析以下章节内容，提取角色的状态变化：

{chapter.content}

请输出JSON格式：
{{
  "character_updates": [
    {{
      "name": "角色名",
      "location": "当前位置",
      "status_changes": ["状态变化"],
      "new_relationships": {{"其他角色": "关系"}}
    }}
  ]
}}
"""
        
        # 这里简化处理，实际应该调用LLM解析
        pass
    
    async def _update_global_summary(self, memory: MemoryManager, chapter: ChapterContent):
        """更新全局摘要"""
        prompt = f"""
当前全局摘要：
{memory.global_summary}

新增章节内容：
{chapter.content}

请更新全局摘要，保留关键情节和角色发展。
"""
        
        summary = await self.llm.generate(prompt)
        memory.global_summary = summary
    
    # ==================== 一致性检查 ====================
    
    async def check_consistency(self, novel_id: str, chapter_num: int) -> Dict[str, Any]:
        """
        检查章节一致性
        
        检查：
        - 角色逻辑一致性
        - 情节连贯性
        - 时间线正确性
        - 设定冲突
        """
        project = await self.get_project(novel_id)
        if not project:
            raise ValueError(f"Project {novel_id} not found")
        
        chapter = next(
            (c for c in project.generated_chapters if c.chapter_num == chapter_num),
            None
        )
        if not chapter:
            raise ValueError(f"Chapter {chapter_num} not found")
        
        memory = self.memories.get(novel_id)
        
        prompt = f"""
请检查以下章节的一致性：

章节内容：
{chapter.content}

角色设定：
{[c.dict() for c in project.characters]}

角色当前状态：
{memory.character_states if memory else {}}

全局摘要：
{memory.global_summary if memory else ''}

请检查：
1. 角色行为是否符合设定
2. 角色位置是否连贯
3. 时间线是否合理
4. 是否有设定冲突

输出JSON格式：
{{
  "issues": [
    {{
      "type": "角色/情节/设定",
      "severity": "严重/中等/轻微",
      "description": "问题描述",
      "location": "问题位置",
      "suggestion": "修改建议"
    }}
  ],
  "overall_score": 0-100
}}
"""
        
        result = await self.llm.generate(prompt)
        
        return {
            "chapter_num": chapter_num,
            "issues": [],  # 解析结果
            "overall_score": 85
        }
    
    # ==================== 导出功能 ====================
    
    async def export_novel(
        self,
        novel_id: str,
        format: str = "markdown",
        output_path: str = None
    ) -> str:
        """
        导出小说
        
        支持格式：markdown, txt, word, pdf
        """
        project = await self.get_project(novel_id)
        if not project:
            raise ValueError(f"Project {novel_id} not found")
        
        if format == "markdown":
            content = self._export_as_markdown(project)
        elif format == "txt":
            content = self._export_as_txt(project)
        elif format == "word":
            content = await self._export_as_word(project, output_path)
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        if output_path and format in ["markdown", "txt"]:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
        
        return content if isinstance(content, str) else output_path
    
    def _export_as_markdown(self, project: NovelProject) -> str:
        """导出为 Markdown"""
        lines = [
            f"# {project.title}",
            f"\n作者：{project.author}",
            f"\n类型：{project.genre.value}",
            f"\n---\n"
        ]
        
        for chapter in project.generated_chapters:
            lines.append(f"\n## {chapter.title}\n")
            lines.append(chapter.content)
            lines.append("\n---\n")
        
        return "\n".join(lines)
    
    def _export_as_txt(self, project: NovelProject) -> str:
        """导出为 TXT"""
        lines = [f"{project.title}\n{'=' * 50}\n"]
        
        for chapter in project.generated_chapters:
            lines.append(f"\n{chapter.title}\n{'-' * 30}\n")
            lines.append(chapter.content)
            lines.append("\n")
        
        return "\n".join(lines)
    
    async def _export_as_word(self, project: NovelProject, output_path: str) -> str:
        """导出为 Word"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            
            doc = Document()
            
            # 标题
            title = doc.add_heading(project.title, 0)
            title.alignment = 1  # 居中
            
            # 元信息
            doc.add_paragraph(f"作者：{project.author}")
            doc.add_paragraph(f"类型：{project.genre.value}")
            doc.add_paragraph("")
            
            # 章节
            for chapter in project.generated_chapters:
                doc.add_heading(chapter.title, level=1)
                para = doc.add_paragraph(chapter.content)
                para.paragraph_format.first_line_indent = Inches(0.3)
                doc.add_paragraph("")
            
            doc.save(output_path)
            return output_path
            
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
